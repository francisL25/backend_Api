from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from pathlib import Path
import subprocess
from ..services.database import get_funcionario_cargo
from ..services.neural_model import get_completion

def guardar_documento(doc, nombre_archivo):
    carpeta_descargas = Path.home() / "Downloads"
    ruta_completa = carpeta_descargas / nombre_archivo
    if ruta_completa.exists():
        contador = 1
        while True:
            nuevo_nombre = f"{nombre_archivo.split('.')[0]}_{contador}.{nombre_archivo.split('.')[1]}"
            ruta_completa_nueva = carpeta_descargas / nuevo_nombre
            if not ruta_completa_nueva.exists():
                ruta_completa = ruta_completa_nueva
                break
            contador += 1
    doc.save(str(ruta_completa))
    if ruta_completa.suffix.lower() == ".docx":
        subprocess.Popen(["start", "", str(ruta_completa)], shell=True)
    return str(ruta_completa)

def crear_nota_interna(tipoDoc, destino, cleaned_suggestions, emisor, referencia, prompt):
    doc = Document(f'app/doc/{tipoDoc}.docx')
    nombres = ''.join([palabra[0] for palabra in destino.split()]) + "/"
    cleaned_suggestions = cleaned_suggestions[::-1]
    if cleaned_suggestions:
        for suggestion in cleaned_suggestions:
            iniciales = ''.join([palabra[0] for palabra in suggestion.split()])
            nombres += iniciales + "/"
    nombres += emisor
    hoja_ruta = ["H.R.:", nombres, "c.c. Archivo"]

    vector_destino = get_funcionario_cargo(destino)
    doc = agregar_texto_a_celda(doc, 0, 2, vector_destino, True)

    if cleaned_suggestions:
        for i, suggestion in enumerate(cleaned_suggestions):
            vector = get_funcionario_cargo(suggestion)
            doc = agregar_texto_a_celda(doc, 1, 2, vector, i == len(cleaned_suggestions) - 1)
    else:
        doc = eliminar_fila_tabla(doc, 0, 2)

    indice_fila = 2 if cleaned_suggestions else 1
    vector_emisor = get_funcionario_cargo(emisor)
    doc = agregar_texto_a_celda(doc, indice_fila, 2, vector_emisor, True)
    vector_referencia = [referencia]
    doc = agregar_texto_a_celda(doc, indice_fila + 1, 2, vector_referencia, True)

    doc.save('./app/doc/tu_documento_modificado.docx')
    doc1 = Document('./app/doc/tu_documento_modificado.docx')
    par = prompt if prompt == "Introduzca su texto" else get_completion(prompt)

    for parra in par.split("\n\n"):
        p = doc1.add_paragraph()
        run = p.add_run(parra)
        run.font.name = 'Century Gothic'
        run.font.size = Pt(10)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    for parra in hoja_ruta:
        p = doc1.add_paragraph()
        run = p.add_run(parra)
        run.font.name = 'Century Gothic'
        run.font.size = Pt(7)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p.paragraph_format.space_after = Pt(0)

    return guardar_documento(doc1, f"{tipoDoc}.docx")

def crear_nota_externa(tipoDoc, destino, referencia, prompt):
    doc = Document(f'app/doc/{tipoDoc}.docx')
    hoja_ruta = ["H.R.:", "Iniciales/...", "c.c. Archivo"]
    referencia = "REF.: " + referencia
    cargo = get_funcionario_cargo(destino)['cargo']
    buscar_reemplazar_texto(doc, "Destino", destino, 'Century Gothic', Pt(10))
    buscar_reemplazar_texto(doc, "CARGO_DESTINO", cargo, 'Century Gothic', Pt(10))
    doc = agregar_text_a_celda(doc, 0, 1, referencia)
    doc.save('./app/doc/tu_documento_modificado.docx')
    doc1 = Document('./app/doc/tu_documento_modificado.docx')
    par = prompt if prompt == "Introduzca su texto" else get_completion(prompt)

    for parra in par.split("\n\n"):
        p = doc1.add_paragraph()
        run = p.add_run(parra)
        run.font.name = 'Century Gothic'
        run.font.size = Pt(10)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    for parra in hoja_ruta:
        p = doc1.add_paragraph()
        run = p.add_run(parra)
        run.font.name = 'Century Gothic'
        run.font.size = Pt(7)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p.paragraph_format.space_after = Pt(0)

    return guardar_documento(doc1, f"{tipoDoc}.docx")

def eliminar_fila_tabla(doc, tabla_index, fila_index):
    tabla = doc.tables[tabla_index]
    tabla._element.remove(tabla._element[1 + fila_index])
    return doc

def buscar_reemplazar_texto(doc, old_text, new_text, font_name, font_size):
    sw = old_text == "CARGO_DESTINO"
    for paragraph in doc.paragraphs:
        if old_text in paragraph.text:
            for run in paragraph.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)
                    run.font.name = font_name
                    run.font.size = font_size
                    if sw:
                        run.font.bold = True

def agregar_text_a_celda(doc, fila, columna, referencia):
    tabla = doc.tables[0]
    celda = tabla.cell(fila, columna)
    p = celda.paragraphs[-1] if celda.paragraphs else celda.add_paragraph()
    run = p.add_run(referencia)
    run.font.name = 'Century Gothic'
    run.font.bold = True
    return doc

def agregar_texto_a_celda(doc, fila, columna, vector_cadenas, estado_final):
    texto = extraer_texto_celda(doc, 0, fila, 0)
    tabla = doc.tables[0]
    celda = tabla.cell(fila, columna)
    p = celda.paragraphs[-1] if celda.paragraphs else celda.add_paragraph()

    for i, parra in enumerate(vector_cadenas):
        run = p.add_run(parra)
        run.font.name = 'Century Gothic'
        run.font.size = Pt(10)
        if texto == "REF.":
            run.font.bold = True
            run.text = run.text.upper()
        if i == 1:
            run.font.bold = True
            run.text = run.text.upper()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p.paragraph_format.space_after = Pt(0)
        if estado_final:
            if i < len(vector_cadenas) - 1:
                p = celda.add_paragraph()
        else:
            p = celda.add_paragraph()
    return doc

def extraer_texto_celda(doc, tabla_index, fila_index, columna_index):
    tabla = doc.tables[tabla_index]
    celda = tabla.cell(fila_index, columna_index)
    return '\n'.join(paragraph.text for paragraph in celda.paragraphs)